from typing import Union, Optional, List

import core.comm
import core.ctrl
from core.pp.elements import *
from core.pp.properties import *

import docx
from docx.document import Document as DocxDoc
from docx.section import Section as DocxSxn
from docx.text.run import Run as DocxRun, Font as DocxFont
from docx.styles.style import _CharacterStyle as DocxCharS, _ParagraphStyle as DocxParaS
from docx.shared import Length as DocxLen
from docx.text.paragraph import Paragraph as DocxPara
from docx.text.parfmt import ParagraphFormat as DocxParFmt
from docx.table import Table as DocxTable, _Row as DocxRow, _Cell as DocxCell

docx_w: str = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
docx_m: str = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'


def get_size(docx_size: Union[DocxLen, float], unit: str) -> Size:
    ans: Size
    if type(docx_size) == float:
        ans = docx_size
    else:
        ans = Length(0, unit)
        if unit == 'cm':
            ans.value = docx_size.cm
        elif unit == 'mm':
            ans.value = docx_size.mm
        elif unit == 'in':
            ans.value = docx_size.inches
        elif unit == 'pt':
            ans.value = docx_size.pt
        else:
            raise UnitError(unit)

    return ans


def get_text_prop_from_docx_font(docx_font: DocxFont) -> TextProp:
    # unit
    # weight
    elem_spacing = None if docx_font.element.rPr is None else docx_font.element.rPr.find(docx_w + 'spacing')
    return TextProp(
        font=Font(
            name=docx_font.name,
            name_asia=None if docx_font.element.rPr is None or docx_font.element.rPr.rFonts is None else docx_font.element.rPr.rFonts.get(
                docx_w + 'eastAsia'),
            size=None if docx_font.size is None else get_size(docx_font.size, 'pt'),
            weight=None if docx_font.bold is None else 'bold' if docx_font.bold is True else 'normal',
        ),
        letter_spacing=None if elem_spacing is None else Length(float(elem_spacing.get(docx_w + 'val')) / 20, 'pt'),
    )


def get_text_prop(docx_char_s: DocxCharS) -> TextProp:
    ans = None if docx_char_s.font is None else get_text_prop_from_docx_font(docx_char_s.font)
    if docx_char_s.base_style is not None:
        ans = core.comm.inherent(ans, get_text_prop(docx_char_s.base_style))
    return ans


def get_span(docx_run: DocxRun) -> Span:
    prop = TextProp()
    if docx_run.font is not None:
        prop = get_text_prop_from_docx_font(docx_run.font)
    if docx_run.style is not None:
        prop = core.comm.inherent(prop, get_text_prop(docx_run.style))

    ans = Span(docx_run.text)
    ans.prop = prop
    return ans


def get_para_prop_from_docx_par_fmt(docx_par_fmt: DocxParFmt) -> ParaProp:
    align: Optional[str] = None
    if docx_par_fmt.alignment is not None:
        if docx_par_fmt.alignment == 0:
            align = 'left'
        elif docx_par_fmt.alignment == 2:
            align = 'right'
        elif docx_par_fmt.alignment == 1:
            align = 'center'
        elif docx_par_fmt.alignment == 3:
            align = 'justify'
        else:
            raise AlignError(docx_par_fmt.alignment)
    return ParaProp(
        align=align,
        line_height=None if docx_par_fmt.line_spacing is None else get_size(docx_par_fmt.line_spacing, 'pt'),
    )


def get_para_prop(docx_para_s: DocxParaS) -> ParaProp:
    ans = None if docx_para_s.paragraph_format is None else get_para_prop_from_docx_par_fmt(
        docx_para_s.paragraph_format)
    if docx_para_s.base_style is not None:
        ans = core.comm.inherent(ans, get_para_prop(docx_para_s.base_style))
    return ans


def get_para(docx_para: DocxPara):
    prop = ParaProp()
    if docx_para.paragraph_format is not None:
        prop = get_para_prop_from_docx_par_fmt(docx_para.paragraph_format)
    if docx_para.style is not None:
        prop = core.comm.inherent(prop, get_para_prop(docx_para.style))

    text_prop = get_text_prop(docx_para.style)

    if len(docx_para._p) == 0:
        return EmptyPara(prop=prop)

    last_is_span: bool = False
    ans = Para(
        prop=prop,
    )
    n_run = 0
    for i in range(0, len(docx_para._p)):
        if docx_para._p[i].tag == docx_m + 'oMath':
            ans.append(Frame(anchor_type='as-char', obj=Math()))
            last_is_span = False
        elif docx_para._p[i].tag == docx_w + 'r':
            if docx_para._p[i].find(docx_w + 'drawing') is not None:
                if docx_para._p[i].find(docx_w + 'drawing').find(docx_w + 'inline') is not None:
                    ans.append(Frame(anchor_type='as-char', obj=Image()))
                else:
                    ans.append(Frame(anchor_type='other', obj=Image()))
                last_is_span = False
            else:
                docx_run = docx_para.runs[n_run]
                n_run += 1
                if docx_run.text == '':
                    continue
                span = get_span(docx_run)
                span.prop = core.comm.inherent(span.prop, text_prop)
                # span.repair()
                if last_is_span and span.prop == ans[-1].prop:
                    new_end = Span(ans[-1] + span)
                    new_end.prop = ans[-1].prop
                    ans[-1] = new_end
                else:
                    ans.append(span)
                last_is_span = True
    return ans


def get_table(docx_table: DocxTable) -> Table:
    align: Optional[str] = None
    if docx_table.alignment is not None:
        if docx_table.alignment == 0:
            align = 'left'
        elif docx_table.alignment == 2:
            align = 'right'
        elif docx_table.alignment == 1:
            align = 'center'
        elif docx_table.alignment == 3:
            align = 'justify'
        else:
            raise AlignError(docx_table.alignment)
    ans = Table(
        prop=ParaProp(
            align=align,
        ),
    )
    docx_row: DocxRow
    for docx_row in docx_table.rows:
        ans.append(Row())
        docx_cell: DocxCell
        for docx_cell in docx_row.cells:
            ans[-1].append(Cell())
            p: DocxPara
            for p in docx_cell.paragraphs:
                ans[-1][-1].append(get_para(p))
    return ans


def get_master_page(docx_sxn: DocxSxn) -> MasterPage:
    ans = MasterPage(
        page_layout=PageLayout(
            prop=PageProp(
                margin_bottom=get_size(docx_sxn.bottom_margin, 'cm') - get_size(docx_sxn.footer_distance, 'cm'),
                margin_left=get_size(docx_sxn.left_margin, 'cm'),
                margin_right=get_size(docx_sxn.right_margin, 'cm'),
                margin_top=get_size(docx_sxn.top_margin, 'cm') - get_size(docx_sxn.header_distance, 'cm'),

                page_height=get_size(docx_sxn.page_height, 'cm'),
                page_width=get_size(docx_sxn.page_width, 'cm'),
            ),
            header_style=HeaderStyle(
                margin_top=get_size(docx_sxn.header_distance, 'cm'),
            ),
            footer_style=FooterStyle(
                margin_bottom=get_size(docx_sxn.footer_distance, 'cm'),
            )
        ),
        header=Header(),
        header_first=Header(),
        footer=Footer(),
        footer_first=Footer(),
    )
    for p in docx_sxn.header.paragraphs:
        ans.header.append(get_para(p))
    if docx_sxn.different_first_page_header_footer:
        for p in docx_sxn.first_page_header.paragraphs:
            ans.header_first.append(get_para(p))
    else:
        ans.header_first = ans.header
    for p in docx_sxn.footer.paragraphs:
        ans.footer.append(get_para(p))
    if docx_sxn.different_first_page_header_footer:
        for p in docx_sxn.first_page_footer.paragraphs:
            ans.footer_first.append(get_para(p))

    return ans


def get_doc(docx_doc: DocxDoc) -> Doc:
    ans = Doc()
    section: List[Union[Para, Table, Break]] = list()
    n_master_page = 0
    master_pages: List[MasterPage] = list()
    for docx_sxn in docx_doc.sections:
        master_pages.append(get_master_page(docx_sxn))
    n_para = 0
    n_table = 0
    # elements = list()
    # for e in docx_doc.element[0]:
    #     if e.tag == docx_w + 'sdt':
    #         if e.find(docx_w + 'sdtContent') is not None:
    #             for ee in e.find(docx_w + 'sdtContent'):
    #                 elements.append(ee)
    #     else:
    #         elements.append(e)
    for e in docx_doc.element[0]:
        if e.tag == docx_w + 'p':
            docx_para = docx_doc.paragraphs[n_para]
            n_para += 1
            # print(n_para)

            if len(
                    docx_para.runs) == 0 and docx_para._p.pPr is not None and docx_para._p.pPr.get_or_add_sectPr() is not None and len(
                    docx_para._p.pPr.get_or_add_sectPr()) > 0:
                ans.append(master_pages[n_master_page])
                n_master_page += 1
                for child in section:
                    ans.append(child)
                section = list()
            else:
                section.append(get_para(docx_para))
                if len(docx_para.runs) > 0:
                    for docx_br in docx_para.runs[0].element.br_lst:
                        if docx_br.get(docx_w + 'type') == 'page':
                            section.append(Break('page'))
                if docx_para._p.pPr is not None and docx_para._p.pPr.get_or_add_sectPr() is not None and len(
                        docx_para._p.pPr.get_or_add_sectPr()) > 0:
                    ans.append(master_pages[n_master_page])
                    n_master_page += 1
                    for child in section:
                        ans.append(child)
                    section = list()


        elif e.tag == docx_w + 'tbl':
            docx_table = docx_doc.tables[n_table]
            n_table += 1
            section.append(get_table(docx_table))
        elif e.tag == docx_w + 'sectPr' or e.tag == docx_w + 'sdt':
            ans.append(master_pages[n_master_page])
            n_master_page += 1
            for child in section:
                ans.append(child)
            section = list()
    return ans


def main(filename: str) -> Doc:
    docx_doc: DocxDoc = docx.Document(filename)
    return get_doc(docx_doc)
